"""
Word Document Extractor
Author: Sarath

Extracts text content from Word documents (.docx) using python-docx.
"""

from pathlib import Path
from typing import Optional, List


class WordExtractor:
    """Extracts content from Word documents."""

    def __init__(self):
        """Initialize Word extractor."""
        try:
            import docx
            self._docx = docx
        except ImportError:
            self._docx = None

    def extract(self, file_path: str) -> 'ExtractedContent':
        """
        Extract text content from a Word document.

        Args:
            file_path: Path to the .docx file

        Returns:
            ExtractedContent object
        """
        from utils.input_router import ExtractedContent

        if self._docx is None:
            raise ImportError("python-docx is required for Word extraction. Install with: pip install python-docx")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract text from Word document
        doc = self._docx.Document(str(path))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = '\n\n'.join(paragraphs)

        # Try to extract title
        title = self._extract_title(doc, path.stem)

        return ExtractedContent(
            text=full_text,
            source_type='word',
            source_path=str(path),
            title=title,
            metadata={
                'paragraph_count': len(paragraphs),
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'core_properties': self._get_core_properties(doc),
            }
        )

    def _extract_title(self, doc, fallback: str) -> str:
        """Extract title from document properties or content."""
        # Try core properties
        try:
            if doc.core_properties.title:
                return doc.core_properties.title
        except Exception:
            pass

        # Try first paragraph if it's a heading style
        for para in doc.paragraphs[:5]:
            if para.style and 'Heading' in para.style.name:
                return para.text.strip()
            if para.text.strip() and len(para.text.strip()) < 100:
                return para.text.strip()

        return fallback.replace('_', ' ').title()

    def _get_core_properties(self, doc) -> dict:
        """Extract core document properties."""
        try:
            props = doc.core_properties
            return {
                'author': props.author,
                'created': str(props.created) if props.created else None,
                'modified': str(props.modified) if props.modified else None,
                'subject': props.subject,
                'keywords': props.keywords,
            }
        except Exception:
            return {}
